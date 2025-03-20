from wordcloud import WordCloud, STOPWORDS
import matplotlib.pyplot as plt
from docx import Document

# Função para extrair texto de um arquivo .docx
def extraindo_texto(docx_path):
    doc = Document(docx_path)
    return "\n".join([item.text for item in doc.paragraphs])

# Caminho do arquivo .docx
docx_path = "gti1.docx"
text = extraindo_texto(docx_path)

# Criando conjunto de stopwords padrão
stopwords = set(STOPWORDS)

# Adicionando novas stopwords do arquivo externo
novas_palavras = set()

try:
    with open(r"C:\oficina\stopwords.txt", 'r', encoding="utf8") as arquivo:
        novas_palavras.update(arquivo.read().split())  # Adiciona todas as palavras separadas por espaço/linha

except FileNotFoundError:
    print("Erro: Arquivo de stopwords não encontrado!")

# Atualizando o conjunto de stopwords
new_stopwords = stopwords.union(novas_palavras)

# Gerando a WordCloud com as stopwords corretas
wordCloud = WordCloud(width=800, height=400, background_color="lightpink", stopwords=new_stopwords).generate(text)

# Exibindo a nuvem de palavras
plt.figure(figsize=(10, 5))
plt.imshow(wordCloud, interpolation="bilinear")
plt.axis("off")
plt.show()


# from wordcloud import WordCloud, STOPWORDS
# import matplotlib.pyplot as plt 
# from docx import Document 


# def extraindo_texto(docx_path):
#     doc = Document(docx_path)
#     texto = ""
#     for item in doc.paragraphs:
#         texto += item.text + "\n" 

#     return texto

# docx_path = 'gti1.docx'

# text = extraindo_texto(docx_path)

# stopwords = set(STOPWORDS)
# novas_palavras = []

# with open(r"C:\\oficina\\stopwords", 'r', encoding="utf8") as item:


#   [novas_palavras.append(word) for linha in item for word in linha.split()]


# new_stopwords = stopwords.union(novas_palavras)

# wordCloud =  WordCloud(width=800, height = 400, background_color="lightpink").generate(text)

# plt.figure(figsize=(10,5))

# plt.imshow(wordCloud, interpolation="bilinear")

# plt.axis("off")

# plt.show()